"""
Script to generate comprehensive DOCX documentation for AI Ingredient Intelligence project
This document will be used as RAG context for the chatbot
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_bullet_point(doc, text, level=0):
    """Add bullet point"""
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    return p

def add_numbered_point(doc, text):
    """Add numbered point"""
    p = doc.add_paragraph(text, style='List Number')
    return p

def create_rag_document():
    """Create comprehensive RAG document for chatbot and landing page"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('AI Ingredient Intelligence Platform - Feature Overview', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Guide to Features, Capabilities, and User Benefits')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.italic = True
    subtitle_format.size = Pt(14)
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        "1. Platform Overview",
        "2. Core Features & Capabilities",
        "3. User Benefits",
        "4. How It Works",
        "5. Use Cases",
        "6. Technical Integration (For Developers)"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Platform Overview
    doc.add_heading('1. Platform Overview', 1)
    doc.add_paragraph(
        'SkinBB AI Ingredient Intelligence is India\'s First Cosmetic Formulation Platform. '
        'It empowers cosmetic formulators, researchers, and manufacturers with AI-powered tools '
        'to analyze, decode, create, and compare cosmetic formulations with unprecedented accuracy and speed.'
    )
    
    doc.add_heading('1.1 What Makes Us Unique', 2)
    unique_features = [
        'First-of-its-kind platform in India for cosmetic formulation intelligence',
        'AI-powered ingredient analysis with regulatory compliance checking',
        'Automatic ingredient extraction from product URLs (Amazon, Nykaa, Flipkart, etc.)',
        'Comprehensive market research capabilities',
        'Professional report generation for stakeholders',
        'Real-time BIS (Bureau of Indian Standards) compliance checking',
        'CAS API integration for accurate ingredient matching',
        'User-friendly interface with intelligent chatbot assistance'
    ]
    for feature in unique_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('1.1 Key Capabilities', 2)
    capabilities = [
        'Ingredient Analysis and Decoding',
        'Formulation Creation and Management',
        'Market Research and Product Discovery',
        'Product Comparison',
        'Compliance Checking (BIS Standards)',
        'Professional Report Generation (PPT)',
        'CAS API Integration for Ingredient Synonyms',
        'URL-based Ingredient Extraction',
        'History Management for Decode and Compare operations'
    ]
    for cap in capabilities:
        add_bullet_point(doc, cap)
    
    # 2. Core Features & Capabilities
    doc.add_heading('2. Core Features & Capabilities', 1)
    
    # 2.1 Decode Formulations
    doc.add_heading('2.1 Decode Formulations', 2)
    doc.add_paragraph(
        'Transform any cosmetic product into a comprehensive ingredient analysis. Simply provide a product URL '
        'or paste an ingredient list, and our AI-powered system will decode every component, providing '
        'detailed insights about each ingredient\'s function, safety, and regulatory status.'
    )
    
    doc.add_heading('Key Capabilities', 3)
    decode_capabilities = [
        'Smart URL Extraction: Automatically extract ingredients from product pages on Amazon, Nykaa, Flipkart, and other e-commerce sites',
        'Flexible Input: Accept ingredient lists in any format - paste directly, upload, or extract from URLs',
        'Intelligent Parsing: Handles various separators and formats automatically',
        'Comprehensive Analysis: Get detailed breakdown of each ingredient including category, function, and safety information',
        'Regulatory Compliance: Automatic BIS (Bureau of Indian Standards) compliance checking with exact limits and cautions',
        'Ingredient Matching: Advanced matching system that handles variations, synonyms, and spelling differences',
        'History Tracking: Save and revisit your previous analyses anytime',
        'Export Options: Generate professional reports for documentation and sharing'
    ]
    for capability in decode_capabilities:
        add_bullet_point(doc, capability)
    
    doc.add_heading('What You Get', 3)
    decode_output = [
        'Complete ingredient breakdown with categorization (Active vs Excipient)',
        'Detailed function and notes for each ingredient',
        'BIS regulatory cautions with exact concentration limits',
        'Ingredient classification (Branded vs General INCI)',
        'Processing time and accuracy metrics',
        'Exportable analysis results'
    ]
    for output in decode_output:
        add_bullet_point(doc, output)
    
    # 2.2 Create Formulations
    doc.add_heading('2.2 Create Formulations', 2)
    doc.add_paragraph(
        'Build new cosmetic formulations from scratch with intelligent ingredient management and '
        'real-time compliance checking. Our platform guides you through the formulation process, '
        'ensuring regulatory compliance and optimal ingredient combinations.'
    )
    
    doc.add_heading('Key Capabilities', 3)
    create_features = [
        'Intuitive Ingredient Management: Easily add, remove, and organize ingredients in your formulation',
        'Real-time Compliance Checking: Get instant feedback on BIS standards compliance as you build',
        'Smart Recommendations: Receive pH range recommendations based on your ingredient selection',
        'Benefit Analysis: Input your desired product benefits and get AI-powered analysis on achievability',
        'Professional Reports: Generate comprehensive reports for stakeholders, investors, or regulatory submissions',
        'Ingredient Validation: Automatic validation of ingredient names and compatibility',
        'Formulation History: Save and manage multiple formulation versions'
    ]
    for feature in create_features:
        add_bullet_point(doc, feature)
    
    # 2.3 Market Research
    doc.add_heading('2.3 Market Research', 2)
    doc.add_paragraph(
        'Discover competitive products and market opportunities with our powerful market research tool. '
        'Find products with similar ingredients, analyze market trends, and identify gaps in the market.'
    )
    
    doc.add_heading('Key Capabilities', 3)
    market_features = [
        'Multiple Search Methods: Search by product URL, ingredient list, product name, or individual ingredients',
        'Smart Matching: Advanced algorithm finds products with similar or matching ingredients',
        'Comprehensive Results: View product details including images, prices, brands, and descriptions',
        'Match Analysis: See exactly which ingredients matched and get match percentage scores',
        'Quick Search: Fast results with processing time under 1 second for most queries',
        'Visual Product Cards: Browse results in an intuitive card-based interface',
        'Filter & Sort: Results sorted by relevance and match quality'
    ]
    for feature in market_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('Search Options', 3)
    search_options = [
        'By Product URL: Paste any product page URL and find similar products',
        'By Ingredient List: Provide your ingredient list to find matching products',
        'By Product Name: Search our database by product name',
        'By Single Ingredient: Find all products containing a specific ingredient (e.g., "Niacinamide")'
    ]
    for option in search_options:
        add_bullet_point(doc, option)
    
    # 2.4 Compare Products
    doc.add_heading('2.4 Compare Products', 2)
    doc.add_paragraph(
        'Make informed decisions by comparing any two cosmetic products side-by-side. Analyze ingredient '
        'overlaps, identify unique components, and understand the differences between competing products.'
    )
    
    doc.add_heading('Key Capabilities', 3)
    compare_features = [
        'Flexible Input: Compare products using URLs, ingredient lists, or a combination of both',
        'Side-by-Side Analysis: Visual comparison showing both products simultaneously',
        'Common Ingredients: Instantly identify shared ingredients between products',
        'Unique Ingredients: Highlight ingredients that appear in only one product',
        'Complete Analysis: Full ingredient breakdown for both products including functions and safety',
        'Save Comparisons: Store comparison results for future reference',
        'Export Results: Download comparison data for reports and presentations'
    ]
    for feature in compare_features:
        add_bullet_point(doc, feature)
    
    # 2.5 Professional Reports
    doc.add_heading('2.5 Professional Reports', 2)
    doc.add_paragraph(
        'Generate comprehensive, professional formulation analysis reports perfect for stakeholders, '
        'regulatory submissions, or internal documentation. Our AI-powered reports provide detailed '
        'insights in an easy-to-understand format.'
    )
    
    doc.add_heading('Report Contents', 3)
    report_sections = [
        'Complete INCI List: All ingredients in your formulation',
        'Detailed Analysis: Ingredient category, functions, and notes for each component',
        'BIS Compliance: Regulatory cautions with exact concentration limits and requirements',
        'Compliance Panel: Regulation status and requirements checklist',
        'Preservative Efficacy: Analysis of preservative systems and stability',
        'Risk Assessment: Risk factors, levels, and mitigation strategies',
        'Benefit Analysis: Cumulative benefits and supporting evidence',
        'Claim Support: Scientific evidence for product claims',
        'pH Recommendations: Optimal pH range with detailed explanations',
        'Expected Benefits Analysis: Assessment of whether your desired benefits are achievable (if provided)'
    ]
    for section in report_sections:
        add_bullet_point(doc, section)
    
    doc.add_heading('Report Formats', 3)
    report_formats = [
        'PowerPoint Presentations: Professional PPTX files perfect for presentations and meetings',
        'JSON Data: Structured data format for integration with other systems',
        'Web View: Interactive HTML format for easy sharing and viewing'
    ]
    for fmt in report_formats:
        add_bullet_point(doc, fmt)
    
    doc.add_heading('AI-Powered Intelligence', 3)
    doc.add_paragraph(
        'Our reports are generated using advanced AI that provides:'
    )
    ai_capabilities = [
        'Intelligent ingredient categorization (Active vs Excipient)',
        'Detailed function descriptions for each ingredient',
        'Accurate BIS regulatory information with exact limits',
        'Comprehensive compliance checking',
        'Evidence-based benefit analysis',
        'Scientific pH recommendations'
    ]
    for capability in ai_capabilities:
        add_bullet_point(doc, capability)
    
    # 3. User Benefits
    doc.add_heading('3. User Benefits', 1)
    
    doc.add_heading('3.1 For Formulators', 2)
    formulator_benefits = [
        'Save Time: Automate ingredient analysis that would take hours manually',
        'Ensure Compliance: Automatic BIS standards checking prevents regulatory issues',
        'Make Informed Decisions: Compare products and analyze market opportunities',
        'Professional Documentation: Generate reports ready for stakeholders and regulatory bodies',
        'Reduce Errors: AI-powered matching reduces human error in ingredient identification',
        'Stay Competitive: Market research helps identify trends and opportunities'
    ]
    for benefit in formulator_benefits:
        add_bullet_point(doc, benefit)
    
    doc.add_heading('3.2 For Researchers', 2)
    researcher_benefits = [
        'Comprehensive Analysis: Get detailed insights into any cosmetic formulation',
        'Regulatory Intelligence: Access BIS cautions and compliance information instantly',
        'Comparative Studies: Easily compare multiple products side-by-side',
        'Data Export: Export analysis results for further research and documentation',
        'Historical Tracking: Save and revisit previous analyses'
    ]
    for benefit in researcher_benefits:
        add_bullet_point(doc, benefit)
    
    doc.add_heading('3.3 For Manufacturers', 2)
    manufacturer_benefits = [
        'Quality Assurance: Verify ingredient compliance before production',
        'Market Intelligence: Research competitive products and market gaps',
        'Regulatory Compliance: Ensure all formulations meet BIS standards',
        'Professional Reports: Generate documentation for regulatory submissions',
        'Cost Efficiency: Reduce time and resources spent on manual analysis'
    ]
    for benefit in manufacturer_benefits:
        add_bullet_point(doc, benefit)
    
    # 4. How It Works
    doc.add_heading('4. How It Works', 1)
    
    doc.add_heading('4.1 Simple Workflow', 2)
    doc.add_paragraph(
        'Our platform is designed for ease of use. Here\'s how users interact with each feature:'
    )
    
    doc.add_heading('4.2 Decode Formulations Workflow', 2)
    decode_workflow = [
        'Step 1: Choose your input method - paste ingredient list or provide product URL',
        'Step 2: Our AI automatically extracts and parses ingredients',
        'Step 3: System matches ingredients against comprehensive databases',
        'Step 4: Receive detailed analysis with categorization, functions, and BIS cautions',
        'Step 5: Save to history or generate professional report'
    ]
    for step in decode_workflow:
        add_numbered_point(doc, step)
    
    doc.add_heading('4.3 Market Research Workflow', 2)
    market_workflow = [
        'Step 1: Select search type (URL, ingredient list, product name, or single ingredient)',
        'Step 2: Enter your search query',
        'Step 3: System searches comprehensive product database',
        'Step 4: View matching products with detailed information and match statistics',
        'Step 5: Analyze results to identify market opportunities'
    ]
    for step in market_workflow:
        add_numbered_point(doc, step)
    
    doc.add_heading('4.4 Compare Products Workflow', 2)
    compare_workflow = [
        'Step 1: Provide first product (URL or ingredient list)',
        'Step 2: Provide second product (URL or ingredient list)',
        'Step 3: System analyzes both products simultaneously',
        'Step 4: View side-by-side comparison with common and unique ingredients highlighted',
        'Step 5: Save comparison for future reference'
    ]
    for step in compare_workflow:
        add_numbered_point(doc, step)
    
    # 5. Use Cases
    doc.add_heading('5. Use Cases', 1)
    
    doc.add_heading('5.1 Product Development', 2)
    doc.add_paragraph(
        'Use our platform during product development to ensure compliance, analyze competitor products, '
        'and validate ingredient combinations before production.'
    )
    
    doc.add_heading('5.2 Regulatory Compliance', 2)
    doc.add_paragraph(
        'Generate comprehensive compliance reports for regulatory submissions. Our BIS integration '
        'ensures all regulatory requirements are met and documented.'
    )
    
    doc.add_heading('5.3 Market Analysis', 2)
    doc.add_paragraph(
        'Research the market to identify trends, find similar products, and discover opportunities '
        'for new formulations or improvements.'
    )
    
    doc.add_heading('5.4 Quality Control', 2)
    doc.add_paragraph(
        'Verify ingredient lists, check compliance, and ensure product safety before manufacturing '
        'or distribution.'
    )
    
    doc.add_heading('5.5 Education & Training', 2)
    doc.add_paragraph(
        'Use the platform for educational purposes to understand ingredient functions, regulatory '
        'requirements, and formulation best practices.'
    )
    
    # 6. Technical Integration (For Developers)
    doc.add_heading('6. Technical Integration (For Developers)', 1)
    
    doc.add_heading('3.1 Ingredient Analysis Endpoints', 2)
    
    doc.add_heading('3.1.1 POST /api/analyze-inci', 3)
    doc.add_paragraph('Analyze ingredients from a JSON payload.')
    doc.add_paragraph('Request Body:', style='Heading 4')
    doc.add_paragraph('{ "inci_names": ["ingredient1", "ingredient2", ...] or "ingredient1, ingredient2" }')
    doc.add_paragraph('Response: AnalyzeInciResponse with matched ingredients, general ingredients, tags, and BIS cautions')
    
    doc.add_heading('3.1.2 POST /api/analyze-inci-form', 3)
    doc.add_paragraph('Analyze ingredients from form data (multipart/form-data).')
    doc.add_paragraph('Request: Form data with inci_names field')
    doc.add_paragraph('Response: AnalyzeInciResponse')
    
    doc.add_heading('3.1.3 POST /api/analyze-url', 3)
    doc.add_paragraph('Extract and analyze ingredients from a product URL.')
    doc.add_paragraph('Request Body: { "url": "https://example.com/product/..." }')
    doc.add_paragraph('Response: AnalyzeInciResponse with extracted text and ingredients')
    
    doc.add_heading('3.1.4 POST /api/extract-ingredients-from-url', 3)
    doc.add_paragraph('Extract ingredients from URL without full analysis.')
    doc.add_paragraph('Request Body: { "url": "https://example.com/product/..." }')
    doc.add_paragraph('Response: ExtractIngredientsResponse with ingredients list and platform info')
    
    doc.add_heading('3.2 Market Research Endpoints', 2)
    
    doc.add_heading('3.2.1 POST /api/market-research', 3)
    doc.add_paragraph('Find products matching ingredients, name, or URL.')
    doc.add_paragraph('Request Body:')
    doc.add_paragraph('{')
    doc.add_paragraph('  "input_type": "url" | "inci" | "name" | "ingredient",')
    doc.add_paragraph('  "url": "..." (if input_type is "url"),')
    doc.add_paragraph('  "inci": "..." (if input_type is "inci"),')
    doc.add_paragraph('  "name": "..." (if input_type is "name"),')
    doc.add_paragraph('  "ingredient": "..." or ["...", "..."] (if input_type is "ingredient")')
    doc.add_paragraph('}')
    doc.add_paragraph('Response: MarketResearchResponse with matched products, match statistics, and processing time')
    
    doc.add_heading('3.3 Product Comparison Endpoints', 2)
    
    doc.add_heading('3.3.1 POST /api/compare-products', 3)
    doc.add_paragraph('Compare two products side-by-side.')
    doc.add_paragraph('Request Body:')
    doc.add_paragraph('{')
    doc.add_paragraph('  "input1": "...", "input1_type": "url" | "inci",')
    doc.add_paragraph('  "input2": "...", "input2_type": "url" | "inci"')
    doc.add_paragraph('}')
    doc.add_paragraph('Response: CompareProductsResponse with comparison data for both products')
    
    doc.add_heading('3.4 History Management Endpoints', 2)
    
    doc.add_heading('3.4.1 Decode History', 3)
    doc.add_paragraph('GET /api/decode-history: Get user\'s decode history')
    doc.add_paragraph('POST /api/decode-history: Save decode history entry')
    doc.add_paragraph('PUT /api/decode-history/{history_id}: Update decode history entry')
    doc.add_paragraph('DELETE /api/decode-history/{history_id}: Delete decode history entry')
    
    doc.add_heading('3.4.2 Compare History', 3)
    doc.add_paragraph('GET /api/compare-history: Get user\'s compare history')
    doc.add_paragraph('POST /api/compare-history: Save compare history entry')
    doc.add_paragraph('PUT /api/compare-history/{history_id}: Update compare history entry')
    doc.add_paragraph('DELETE /api/compare-history/{history_id}: Delete compare history entry')
    
    doc.add_heading('3.5 Formulation Report Endpoints', 2)
    
    doc.add_heading('3.5.1 POST /api/formulation-report-json', 3)
    doc.add_paragraph('Generate formulation report as JSON.')
    doc.add_paragraph('Request Body:')
    doc.add_paragraph('{')
    doc.add_paragraph('  "inciList": ["ingredient1", "ingredient2", ...],')
    doc.add_paragraph('  "brandedIngredients": [...],')
    doc.add_paragraph('  "notBrandedIngredients": [...],')
    doc.add_paragraph('  "bisCautions": {...},')
    doc.add_paragraph('  "expectedBenefits": "..." (optional)')
    doc.add_paragraph('}')
    doc.add_paragraph('Response: FormulationReportResponse with all report sections')
    
    doc.add_heading('3.5.2 POST /api/formulation-report/ppt', 3)
    doc.add_paragraph('Generate PowerPoint presentation from report JSON.')
    doc.add_paragraph('Request Body: { "reportData": FormulationReportResponse }')
    doc.add_paragraph('Response: PPTX file download')
    
    doc.add_heading('3.6 Health and Info Endpoints', 2)
    
    doc.add_heading('3.6.1 GET /api/server-health', 3)
    doc.add_paragraph('Comprehensive server health check (Chrome, APIs, environment variables)')
    
    doc.add_heading('3.6.2 GET /api/info', 3)
    doc.add_paragraph('API information and available endpoints')
    
    doc.add_heading('3.6.3 GET /health', 3)
    doc.add_paragraph('Basic health check endpoint')
    
    # 4. Frontend Routes and Pages
    doc.add_heading('4. Frontend Routes and Pages', 1)
    
    doc.add_heading('4.1 Formulator Module Routes', 2)
    routes = [
        ('/dashboard', 'Formulator Dashboard - Main landing page with feature cards'),
        ('/formulations/decode', 'Decode Formulations - Analyze existing formulations'),
        ('/formulations/create', 'Create Formulations - Build new formulations'),
        ('/market-research', 'Market Research - Find products with matching ingredients'),
        ('/compare', 'Compare Products - Side-by-side product comparison'),
        ('/account', 'Account Settings - Manage user account and preferences')
    ]
    for route, description in routes:
        p = doc.add_paragraph()
        p.add_run(f'{route}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('4.2 Feature Pages', 2)
    
    doc.add_heading('4.2.1 Decode Formulations Page', 3)
    decode_features = [
        'Input area for INCI list or URL',
        'History panel to view past analyses',
        'Results display with ingredient breakdown',
        'BIS cautions display',
        'Save to history functionality',
        'Export options'
    ]
    for feature in decode_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('4.2.2 Market Research Page', 3)
    market_page_features = [
        'Multiple input modes: INCI, URL, Name, Ingredient',
        'Input mode toggle buttons',
        'Extract ingredients from URL button',
        'Results grid with product cards',
        'Product images, prices, and match statistics',
        'Matched ingredients display',
        'Processing time indicator'
    ]
    for feature in market_page_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('4.2.3 Compare Products Page', 3)
    compare_page_features = [
        'Dual input areas for two products',
        'Input type selection for each product',
        'Side-by-side comparison view',
        'Common ingredients highlighting',
        'Unique ingredients identification',
        'History management',
        'Export comparison results'
    ]
    for feature in compare_page_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('4.2.4 Formulation Report Viewer', 3)
    report_viewer_features = [
        'Interactive report display',
        'Section-by-section navigation',
        'PPT download button',
        'Print functionality',
        'Responsive design for all devices'
    ]
    for feature in report_viewer_features:
        add_bullet_point(doc, feature)
    
    # 5. Data Models and Schemas
    doc.add_heading('5. Data Models and Schemas', 1)
    
    doc.add_heading('5.1 Request Models', 2)
    doc.add_paragraph('AnalyzeInciRequest: { "inci_names": string | List[string] }')
    doc.add_paragraph('ExtractIngredientsRequest: { "url": string }')
    doc.add_paragraph('MarketResearchRequest: { "input_type": string, "url" | "inci" | "name" | "ingredient": ... }')
    doc.add_paragraph('CompareProductsRequest: { "input1": string, "input1_type": string, "input2": string, "input2_type": string }')
    doc.add_paragraph('FormulationReportRequest: { "inciList": List[string], "brandedIngredients": List[string], "notBrandedIngredients": List[string], "bisCautions": Dict, "expectedBenefits": string }')
    
    doc.add_heading('5.2 Response Models', 2)
    doc.add_paragraph('AnalyzeInciResponse: Contains matched ingredients, general ingredients, unable to decode list, tags, BIS cautions, processing time')
    doc.add_paragraph('ExtractIngredientsResponse: Contains ingredients list, extracted text, platform, processing time, source')
    doc.add_paragraph('MarketResearchResponse: Contains products list, extracted ingredients, total matched, processing time, input type')
    doc.add_paragraph('CompareProductsResponse: Contains comparison data for both products, common ingredients, unique ingredients')
    doc.add_paragraph('FormulationReportResponse: Contains all report sections (INCI list, analysis table, compliance panel, etc.)')
    
    # 6. Integration Guide
    doc.add_heading('6. Integration Guide', 1)
    
    doc.add_heading('6.1 URL Scraping', 2)
    doc.add_paragraph(
        'The platform supports automatic ingredient extraction from various e-commerce platforms:'
    )
    supported_platforms = [
        'Amazon',
        'Nykaa',
        'Flipkart',
        'Other e-commerce sites with ingredient information'
    ]
    for platform in supported_platforms:
        add_bullet_point(doc, platform)
    
    doc.add_heading('6.2 CAS API Integration', 2)
    doc.add_paragraph(
        'The platform integrates with CAS API to retrieve ingredient synonyms for better matching. '
        'This helps match ingredients even when they have different names or spellings.'
    )
    
    doc.add_heading('6.3 BIS RAG Integration', 2)
    doc.add_paragraph(
        'BIS (Bureau of Indian Standards) cautions are retrieved using RAG (Retrieval Augmented Generation) '
        'from PDF documents. The system uses vector embeddings and semantic search to find relevant '
        'regulatory information for each ingredient.'
    )
    
    doc.add_heading('6.4 MongoDB Collections', 2)
    collections = [
        'externalProducts: External product database for market research',
        'ingre_branded_ingredients: Branded ingredient database',
        'ingre_inci: General INCI ingredient database',
        'decode_history: User decode history',
        'compare_history: User compare history',
        'ingre_functional_categories: Functional category information',
        'ingre_chemical_classes: Chemical class information'
    ]
    for collection in collections:
        add_bullet_point(doc, collection)
    
    doc.add_paragraph(
        'This section provides technical details for developers who need to integrate with the platform.'
    )
    
    doc.add_heading('6.1 API Endpoints', 2)
    doc.add_paragraph('The platform provides RESTful API endpoints for programmatic access.')
    
    doc.add_heading('6.2 Frontend Routes', 2)
    routes = [
        '/dashboard - Main dashboard with feature overview',
        '/formulations/decode - Decode formulations page',
        '/formulations/create - Create formulations page',
        '/market-research - Market research page',
        '/compare - Compare products page',
        '/account - Account settings page'
    ]
    for route in routes:
        add_bullet_point(doc, route)
    
    doc.add_heading('6.3 Key Integrations', 2)
    integrations = [
        'CAS API: For ingredient synonym matching',
        'BIS RAG: For regulatory compliance information',
        'MongoDB: For ingredient and product databases',
        'OpenAI/Claude: For AI-powered analysis and report generation',
        'Gamma API: For PowerPoint presentation generation'
    ]
    for integration in integrations:
        add_bullet_point(doc, integration)
    
    # Chatbot Integration Section
    doc.add_page_break()
    doc.add_heading('8. Chatbot Integration Guide', 1)
    
    doc.add_heading('8.1 Chatbot Features', 2)
    chatbot_features = [
        'Intent Detection: Automatically detects user intent (decode, create, market research, compare, account)',
        'Conversation Flow: Guides users through feature discovery',
        'Redirect Buttons: Provides direct navigation buttons to relevant features',
        'Platform Information: Explains platform capabilities',
        'Inspiration Questions: Asks about reference products or inspirations'
    ]
    for feature in chatbot_features:
        add_bullet_point(doc, feature)
    
    doc.add_heading('8.2 Chatbot Routes', 2)
    chatbot_routes = [
        'POST /api/chatbot/chat: Main chatbot endpoint for formulator module',
        'POST /api/chat: General chatbot endpoint with RAG'
    ]
    for route in chatbot_routes:
        add_bullet_point(doc, route)
    
    doc.add_heading('8.3 Redirect Button Format', 2)
    doc.add_paragraph('The chatbot can provide redirect buttons with the following format:')
    doc.add_paragraph('{')
    doc.add_paragraph('  "label": "Go to [Feature Name]",')
    doc.add_paragraph('  "route": "/formulations/decode" | "/formulations/create" | "/market-research" | "/compare" | "/account",')
    doc.add_paragraph('  "description": "Feature description"')
    doc.add_paragraph('}')
    
    doc.add_heading('8.4 Intent Detection Keywords', 2)
    intent_keywords = {
        'decode': ['decode', 'analyze', 'break down', 'understand', 'ingredient breakdown'],
        'create': ['create', 'make', 'new formulation', 'build', 'develop'],
        'market research': ['market research', 'find products', 'search', 'research', 'similar products'],
        'compare': ['compare', 'comparison', 'side by side'],
        'account': ['account', 'settings', 'profile', 'preferences']
    }
    for intent, keywords in intent_keywords.items():
        p = doc.add_paragraph()
        p.add_run(f'{intent.title()}: ').bold = True
        p.add_run(', '.join(keywords))
    
    # Save document
    doc.save('AI_Ingredient_Intelligence_Complete_Documentation.docx')
    print("Documentation generated successfully: AI_Ingredient_Intelligence_Complete_Documentation.docx")

if __name__ == "__main__":
    try:
        from docx import Document
        create_rag_document()
    except ImportError:
        print("python-docx library not found. Installing...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        create_rag_document()

